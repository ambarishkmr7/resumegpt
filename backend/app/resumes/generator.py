"""Render ResumeContent to downloadable PDF and DOCX files.

Template styling is driven by app.templates.registry. Adding a template is a
matter of adding a style dict — the renderers read colors/spacing from it.
"""
import io

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    Paragraph, SimpleDocTemplate, Spacer, HRFlowable, ListFlowable, ListItem,
)

from app.schemas import ResumeContent
from app.templates.registry import get_template


# ---------------- PDF ----------------

def render_pdf(content: ResumeContent, template_id: str) -> bytes:
    tpl = get_template(template_id)
    accent = HexColor(tpl["accent"])
    is_serif = tpl.get("font") == "serif"
    centered = tpl.get("layout") == "centered"
    base_font = "Times-Roman" if is_serif else "Helvetica"
    bold_font = "Times-Bold" if is_serif else "Helvetica-Bold"
    name_align = 1 if centered else 0  # 1 = TA_CENTER

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
    )
    ss = getSampleStyleSheet()
    name_style = ParagraphStyle("Name", parent=ss["Title"], fontSize=22,
                                textColor=accent, spaceAfter=2, leading=24,
                                fontName=bold_font, alignment=name_align)
    title_style = ParagraphStyle("Headline", parent=ss["Normal"], fontSize=11,
                                  textColor=HexColor("#444444"), spaceAfter=4,
                                  fontName=base_font, alignment=name_align)
    contact_style = ParagraphStyle("Contact", parent=ss["Normal"], fontSize=9,
                                   textColor=HexColor("#555555"), spaceAfter=8,
                                   fontName=base_font, alignment=name_align)
    section_style = ParagraphStyle("Section", parent=ss["Heading2"], fontSize=12,
                                   textColor=accent, spaceBefore=10, spaceAfter=3,
                                   fontName=bold_font)
    body = ParagraphStyle("Body", parent=ss["Normal"], fontSize=9.5, leading=13,
                          fontName=base_font)
    role_style = ParagraphStyle("Role", parent=ss["Normal"], fontSize=10.5,
                                leading=13, spaceBefore=4, fontName=bold_font)
    meta_style = ParagraphStyle("Meta", parent=ss["Normal"], fontSize=9,
                                textColor=HexColor("#666666"), spaceAfter=2,
                                fontName=base_font)

    flow = []
    c = content.contact
    if c.name:
        flow.append(Paragraph(c.name, name_style))
    if c.title:
        flow.append(Paragraph(c.title, title_style))
    contact_bits = [b for b in [c.email, c.phone, c.location, c.linkedin, c.website] if b]
    if contact_bits:
        flow.append(Paragraph("  •  ".join(contact_bits), contact_style))

    def section(label: str):
        flow.append(Paragraph(label.upper(), section_style))
        flow.append(HRFlowable(width="100%", thickness=1, color=accent, spaceAfter=4))

    if content.summary:
        section("Summary")
        flow.append(Paragraph(content.summary, body))

    if content.experience:
        section("Experience")
        for e in content.experience:
            header = " — ".join([x for x in [e.title, e.company] if x])
            flow.append(Paragraph(header, role_style))
            meta = "  |  ".join([x for x in [e.location, f"{e.start} – {e.end}".strip(" –")] if x])
            if meta:
                flow.append(Paragraph(meta, meta_style))
            if e.bullets:
                flow.append(ListFlowable(
                    [ListItem(Paragraph(b, body), leftIndent=10) for b in e.bullets],
                    bulletType="bullet", start="•", leftIndent=12,
                ))

    if content.projects:
        section("Projects")
        for p in content.projects:
            flow.append(Paragraph(p.name, role_style))
            if p.description:
                flow.append(Paragraph(p.description, body))
            if p.bullets:
                flow.append(ListFlowable(
                    [ListItem(Paragraph(b, body)) for b in p.bullets],
                    bulletType="bullet", start="•", leftIndent=12,
                ))

    if content.education:
        section("Education")
        for ed in content.education:
            header = " — ".join([x for x in [ed.degree, ed.school] if x])
            flow.append(Paragraph(header, role_style))
            meta = "  |  ".join([x for x in [ed.location, f"{ed.start} – {ed.end}".strip(" –")] if x])
            if meta:
                flow.append(Paragraph(meta, meta_style))
            if ed.details:
                flow.append(Paragraph(ed.details, body))

    if content.skills:
        section("Skills")
        flow.append(Paragraph("  •  ".join(content.skills), body))

    if content.certifications:
        section("Certifications")
        flow.append(ListFlowable(
            [ListItem(Paragraph(x, body)) for x in content.certifications],
            bulletType="bullet", start="•", leftIndent=12,
        ))

    if content.languages:
        section("Languages")
        flow.append(Paragraph("  •  ".join(content.languages), body))

    if content.accomplishments:
        section("Accomplishments")
        flow.append(ListFlowable(
            [ListItem(Paragraph(x, body)) for x in content.accomplishments],
            bulletType="bullet", start="•", leftIndent=12,
        ))

    doc.build(flow)
    return buf.getvalue()


# ---------------- DOCX ----------------

def render_docx(content: ResumeContent, template_id: str) -> bytes:
    tpl = get_template(template_id)
    accent_hex = tpl["accent"].lstrip("#")
    accent_rgb = RGBColor(int(accent_hex[0:2], 16), int(accent_hex[2:4], 16), int(accent_hex[4:6], 16))
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.6)
        s.left_margin = s.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    c = content.contact
    if c.name:
        h = doc.add_paragraph()
        run = h.add_run(c.name)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = accent_rgb
    if c.title:
        doc.add_paragraph(c.title)
    contact_bits = [b for b in [c.email, c.phone, c.location, c.linkedin, c.website] if b]
    if contact_bits:
        p = doc.add_paragraph("  •  ".join(contact_bits))
        p.runs[0].font.size = Pt(9)

    def section(label: str):
        p = doc.add_paragraph()
        run = p.add_run(label.upper())
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = accent_rgb
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)

    if content.summary:
        section("Summary")
        doc.add_paragraph(content.summary)

    if content.experience:
        section("Experience")
        for e in content.experience:
            p = doc.add_paragraph()
            r = p.add_run(" — ".join([x for x in [e.title, e.company] if x]))
            r.bold = True
            meta = "  |  ".join([x for x in [e.location, f"{e.start} – {e.end}".strip(" –")] if x])
            if meta:
                mp = doc.add_paragraph(meta)
                mp.runs[0].font.size = Pt(9)
            for b in e.bullets:
                doc.add_paragraph(b, style="List Bullet")

    if content.projects:
        section("Projects")
        for proj in content.projects:
            p = doc.add_paragraph()
            p.add_run(proj.name).bold = True
            if proj.description:
                doc.add_paragraph(proj.description)
            for b in proj.bullets:
                doc.add_paragraph(b, style="List Bullet")

    if content.education:
        section("Education")
        for ed in content.education:
            p = doc.add_paragraph()
            p.add_run(" — ".join([x for x in [ed.degree, ed.school] if x])).bold = True
            meta = "  |  ".join([x for x in [ed.location, f"{ed.start} – {ed.end}".strip(" –")] if x])
            if meta:
                doc.add_paragraph(meta).runs[0].font.size = Pt(9)
            if ed.details:
                doc.add_paragraph(ed.details)

    if content.skills:
        section("Skills")
        doc.add_paragraph("  •  ".join(content.skills))

    if content.certifications:
        section("Certifications")
        for x in content.certifications:
            doc.add_paragraph(x, style="List Bullet")

    if content.languages:
        section("Languages")
        doc.add_paragraph("  •  ".join(content.languages))

    if content.accomplishments:
        section("Accomplishments")
        for x in content.accomplishments:
            doc.add_paragraph(x, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
